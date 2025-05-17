import os
import logging
import re
import fitz  # PyMuPDF
import tempfile
import docx
from docx import Document
from docx.shared import Inches
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    MessageHandler,
    CallbackQueryHandler,
    filters,
    ContextTypes,
)
from PyPDF2 import PdfReader

# --- Логирование ---
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "Привет! Пришли мне PDF — я верну тебе текст и картинки, если они есть."
    )

async def handle_pdf(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("⏳ Обрабатываю файл, это может занять несколько секунд...")

    doc_file = update.message.document
    if doc_file.mime_type != "application/pdf":
        await context.bot.send_message(
            chat_id=update.effective_chat.id,
            text="Это не PDF."
        )
        return

    file_path = f"/tmp/{doc_file.file_name}"
    new_file = await context.bot.get_file(doc_file.file_id)
    await new_file.download_to_drive(file_path)

    # --- Собираем по страницам: текст и картинки ---
    pdf_doc = fitz.open(file_path)
    reader = PdfReader(file_path)
    docx_blocks = []
    sent_hashes = set()
    num_pages = len(pdf_doc)

    for i in range(num_pages):
        # --- текст ---
        page_text = ""
        try:
            page_text = reader.pages[i].extract_text() or ""
            page_text = re.sub(r"(\w)-\n(\w)", r"\1\2", page_text)
            page_text = page_text.replace("\n", " ")
            page_text = re.sub(r" {2,}", " ", page_text).strip()
        except Exception:
            pass
        if page_text:
            docx_blocks.append(("text", page_text))

        # --- картинки ---
        for img in pdf_doc[i].get_images(full=True):
            xref = img[0]
            img_dict = pdf_doc.extract_image(xref)
            img_bytes = img_dict['image']
            ext = img_dict['ext']
            img_hash = hash(img_bytes)
            if img_hash not in sent_hashes:
                docx_blocks.append(("image", (img_bytes, ext)))
                sent_hashes.add(img_hash)

    # --- сохраняем блоки для Word ---
    context.user_data['last_pdf_blocks'] = docx_blocks

    # --- отправка в чат: текст и картинки (только уникальные, в порядке следования) ---
    for block_type, block_val in docx_blocks:
        if block_type == "text":
            for i in range(0, len(block_val), 4096):
                await context.bot.send_message(chat_id=update.effective_chat.id, text=block_val[i:i+4096])
        elif block_type == "image":
            img_bytes, ext = block_val
            await context.bot.send_photo(
                chat_id=update.effective_chat.id,
                photo=img_bytes
            )

    # --- кнопки ---
    keyboard = InlineKeyboardMarkup([
        [InlineKeyboardButton("📥 Скачать в Word", callback_data="download_word")],
        [InlineKeyboardButton("🔄 Загрузить ещё PDF-файл", callback_data="start_over")],
    ])
    await context.bot.send_message(
        chat_id=update.effective_chat.id,
        text="Ваш текст готов! Выберите действие ниже:",
        reply_markup=keyboard,
    )

async def download_word_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    docx_blocks = context.user_data.get('last_pdf_blocks', [])

    if not docx_blocks:
        return await query.edit_message_text("Нет данных для генерации Word.")

    doc = Document()
    for block_type, block_val in docx_blocks:
        if block_type == "text":
            doc.add_paragraph(block_val)
        elif block_type == "image":
            img_bytes, ext = block_val
            with tempfile.NamedTemporaryFile(delete=False, suffix='.' + ext) as tmp_img:
                tmp_img.write(img_bytes)
                tmp_img.flush()
                try:
                    doc.add_picture(tmp_img.name, width=Inches(5))
                except Exception as e:
                    logger.warning(f'Не удалось вставить картинку: {e}')

    out_path = "/tmp/output.docx"
    doc.save(out_path)

    await context.bot.send_document(
        chat_id=update.effective_chat.id,
        document=open(out_path, 'rb'),
        filename='converted.docx'
    )

    keyboard = InlineKeyboardMarkup([
        [InlineKeyboardButton("🔄 Загрузить ещё PDF-файл", callback_data="start_over")],
    ])
    await query.edit_message_reply_markup(reply_markup=keyboard)

async def start_over_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    context.user_data.pop('last_pdf_blocks', None)
    await query.edit_message_reply_markup(None)
    await context.bot.send_message(
        chat_id=update.effective_chat.id,
        text="🔄 Пришлите новый PDF-файл сюда."
    )

def main():
    token = os.getenv('TELEGRAM_BOT_TOKEN')
    if not token:
        logger.error('TELEGRAM_BOT_TOKEN не задан')
        return

    app = (
        ApplicationBuilder()
        .token(token)
        .build()
    )
    app.add_handler(CommandHandler('start', start))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_pdf))
    app.add_handler(CallbackQueryHandler(download_word_callback, pattern='download_word'))
    app.add_handler(CallbackQueryHandler(start_over_callback, pattern='start_over'))

    host = os.getenv('RENDER_EXTERNAL_URL')
    if not host:
        logger.error('RENDER_EXTERNAL_URL не задан')
        return
    port = int(os.getenv('PORT', 5000))
    webhook_url = f"{host}/{token}"

    app.run_webhook(
        listen='0.0.0.0',
        port=port,
        url_path=token,
        webhook_url=webhook_url
    )

if __name__ == '__main__':
    main()
